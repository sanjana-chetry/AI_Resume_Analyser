"""
resume_parser.py
----------------
Extracts raw text from PDF or DOCX resume files.

Libraries needed:
    pip install pdfplumber python-docx
"""

import os
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract all text from a PDF using pdfplumber.
    Raises RuntimeError if no text was found (e.g. scanned image PDF).
    """
    text = ""

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    if not text.strip():
        raise RuntimeError(
            "No text could be extracted from this PDF. "
            "It may be a scanned image. Please upload a text-based PDF."
        )

    return text


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a DOCX file.
    Also extracts text from tables — many resumes use tables for skills layout.
    """
    doc = Document(file_path)
    parts = []

    # Extract paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Extract text from tables (skills are often listed in table cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    if not parts:
        raise RuntimeError(
            "No text found in the DOCX file. "
            "The file may be empty or corrupted."
        )

    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    """
    Extract raw text from a resume file.

    Supports: .pdf, .docx

    Raises
    ------
    ValueError   — unsupported file format
    RuntimeError — extraction failed (empty, corrupted, or scanned file)
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        # Raise so app.py can catch it and show the user a proper error message
        raise ValueError(
            f"Unsupported file format '{ext}'. Please upload a PDF or DOCX file."
        )
